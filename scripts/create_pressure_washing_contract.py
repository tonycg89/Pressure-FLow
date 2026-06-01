from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Pressure Washing Service Agreement.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
LIGHT_FILL = "F2F4F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D0D5DD"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = width
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_field_table(doc):
    fields = [
        ("Business", "[Business legal name]"),
        ("Client", "[Client name]"),
        ("Service Address", "[Service address]"),
        ("Approved Estimate", "[Square estimate number or link]"),
        ("Estimated Price", "[$ amount]"),
        ("Deposit", "[% or $ amount]"),
        ("Scheduled Date", "[Date/time, if known]"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    set_table_width(table, [Inches(1.8), Inches(4.7)])
    header = table.rows[0].cells
    header[0].text = "Field"
    header[1].text = "Details"
    for cell in header:
        set_cell_shading(cell, LIGHT_FILL)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for label, detail in fields:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
    return table


def add_clause(doc, number, title, body, initials=False):
    heading = doc.add_paragraph(style="Heading 2")
    heading.add_run(f"{number}. {title}")
    for paragraph_text in body:
        p = doc.add_paragraph(paragraph_text)
        p.paragraph_format.keep_together = True
    if initials:
        p = doc.add_paragraph()
        run = p.add_run("Client initials required: ________")
        run.bold = True
        run.font.color.rgb = DARK_BLUE
        p.paragraph_format.space_after = Pt(8)


def add_signature_block(doc):
    doc.add_paragraph(style="Heading 1").add_run("Signatures")
    p = doc.add_paragraph(
        "By signing below, the Client and Business agree to the terms of this Service Agreement."
    )
    p.paragraph_format.keep_with_next = True

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    set_table_width(table, [Inches(3.25), Inches(3.25)])

    rows = [
        ("Client Signature", "Business Signature"),
        ("Printed Name: __________________________", "Printed Name: __________________________"),
        ("Date: _________________________________", "Date: _________________________________"),
        ("Email/Phone: __________________________", "Title: _________________________________"),
    ]

    for row, values in zip(table.rows, rows):
        for cell, value in zip(row.cells, values):
            cell.text = value
            if row == table.rows[0]:
                set_cell_shading(cell, LIGHT_FILL)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(32, 33, 36)
    title.paragraph_format.space_after = Pt(6)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = RGBColor(102, 112, 133)
    subtitle.paragraph_format.space_after = Pt(12)

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.color.rgb = BLUE
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.color.rgb = BLUE
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True


def main():
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Pressure Washing Service Agreement")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Draft template for Square Contracts. Review with a qualified professional before customer use.")

    doc.add_paragraph(style="Heading 1").add_run("Project Details")
    add_field_table(doc)

    doc.add_paragraph(style="Heading 1").add_run("Terms and Conditions")

    clauses = [
        (
            "Scope of Work",
            [
                "The services to be performed are limited to the scope of work listed in the approved estimate. Additional areas, stains, treatments, repairs, restoration, stain removal, oxidation removal, rust removal, oil removal, or other services are not included unless agreed to in writing by the Business and may require additional charges."
            ],
            True,
        ),
        (
            "Estimate Approval",
            [
                "The Client acknowledges that the approved estimate identifies the services, pricing, and scope of work for the project. By approving or signing the estimate, the Client authorizes the Business to proceed with the services described in the estimate, subject to the terms of this Contract. Any services, areas, conditions, or work not listed in the approved estimate are excluded unless agreed to in writing by the Business and may require additional charges."
            ],
            False,
        ),
        (
            "Additional Costs",
            [
                "The approved estimate includes only the services and pricing listed in the estimate. Additional costs may apply for extra services, added areas, unusual access requirements, special stain treatments, required materials, parking fees, disposal fees, water access issues, or other costs not included in the approved estimate. Any additional costs must be approved by the Client in writing before they are added to the final invoice."
            ],
            False,
        ),
        (
            "Access and Utilities",
            [
                "The Client is responsible for providing safe and reasonable access to all areas where services are to be performed, including unlocked gates, clear work areas, accessible water sources, and any required parking or entry instructions. The Client is responsible for ensuring that exterior water spigots are working and available unless otherwise agreed in writing. If access, water, parking, or required utilities are unavailable at the scheduled service time, the Business may reschedule the service and the Client may be responsible for a rescheduling fee or forfeiture of deposit as stated in this Contract."
            ],
            True,
        ),
        (
            "Pre-Existing Conditions",
            [
                "The Client acknowledges that surfaces may have pre-existing conditions including oxidation, fading, loose paint, failing stain or sealant, cracked concrete, damaged siding, deteriorated wood, loose mortar, damaged screens, worn seals, leaking windows or doors, rust, efflorescence, algae staining, mold, mildew, or other wear. Cleaning may reveal or make these conditions more visible. The Business is not responsible for pre-existing damage or conditions, or for cosmetic differences caused by age, wear, oxidation, fading, or prior maintenance."
            ],
            True,
        ),
        (
            "Property Condition and Damage Limitations",
            [
                "The Client understands that pressure washing, soft washing, detergents, and cleaning solutions involve water, pressure, and chemicals that may affect surfaces, seals, landscaping, fixtures, paint, coatings, oxidation, stains, loose materials, or pre-existing damage. The Business will use reasonable care while performing the services, but is not responsible for damage caused by pre-existing conditions, defective or deteriorated materials, poor installation, loose paint, oxidation, failing caulking or seals, unsealed doors or windows, water intrusion through existing gaps, improperly protected electrical components, fragile surfaces, or items left in the work area.",
                "The Client is responsible for closing and securing all windows and doors, identifying known leaks or sensitive areas, moving or protecting personal property, vehicles, decor, plants, furniture, and fragile items, and notifying the Business of any surfaces or areas requiring special care before work begins.",
                "If the Business directly causes damage through negligence, the Business's responsibility is limited to the reasonable cost of repair for the directly affected area, and does not include pre-existing damage, indirect damages, lost use, lost profits, replacement of unrelated materials, or cosmetic differences caused by age, wear, fading, oxidation, or prior condition.",
            ],
            True,
        ),
        (
            "Plants, Landscaping, and Outdoor Items",
            [
                "The Client understands that detergents, cleaning solutions, runoff, overspray, and rinsing may contact plants, grass, landscaping, soil, mulch, outdoor furniture, decor, vehicles, or other items near the work area. The Business will use reasonable care to rinse or avoid affected areas when practical, but the Client is responsible for moving, covering, watering, or otherwise protecting sensitive plants, personal property, vehicles, furniture, decor, and fragile outdoor items before service begins.",
                "The Business is not responsible for damage, spotting, discoloration, plant stress, plant loss, or other effects caused by pre-existing plant health, drought, heat, poor drainage, chemicals or treatments previously applied by others, items left in the work area, or the Client's failure to identify sensitive areas before work begins. If the Business directly causes damage through negligence, the Business's responsibility is limited to the reasonable cost of repair or replacement of the directly affected item.",
            ],
            True,
        ),
        (
            "Weather and Rescheduling",
            [
                "The Business may reschedule services due to weather, unsafe conditions, equipment issues, water restrictions, or other conditions that may affect safety or service quality. If services are rescheduled by the Business, the Client's deposit will remain applied to the rescheduled service date."
            ],
            False,
        ),
        (
            "Deposit",
            [
                "A deposit is required before the work is scheduled. The deposit will be applied toward the total project price. The deposit is refundable if the Client cancels at least 48 hours before the scheduled service date. The deposit is non-refundable if the Client cancels within 48 hours of the scheduled service date, is unavailable at the scheduled time, fails to provide required access, or prevents the Business from performing the work as scheduled. If the Business must cancel and cannot reschedule within a reasonable time, the deposit will be refunded."
            ],
            True,
        ),
        (
            "Invoices and Payment",
            [
                "The Client agrees to pay the required deposit before the work is scheduled. After the services are completed, the Business will send a final invoice for the remaining balance. The Client agrees to pay all invoices by the due date stated on the invoice. Unpaid or overdue invoices may result in delayed scheduling, suspension of additional services, collection efforts, or other remedies allowed by law."
            ],
            True,
        ),
        (
            "Payment Methods",
            [
                "Payment may be made through the payment link or invoice provided by the Business, including approved credit card, debit card, or other payment methods supported by Square. The Business may also accept cash, check, or another payment method at its discretion. Payment is not considered received until funds have cleared."
            ],
            False,
        ),
        (
            "Service Acceptance, Corrections, and Refunds",
            [
                "The Client agrees to inspect the completed work promptly after service is performed. If the Client believes an agreed-upon service was not completed, the Client must notify the Business within 24 hours and allow the Business a reasonable opportunity to review and, if appropriate, correct the issue.",
                "Payments are non-refundable after the service has been completed, except where required by law or agreed to in writing by the Business. Refunds will not be issued for pre-existing conditions, stains or discoloration that cannot be fully removed, oxidation, wear, damage not caused by the Business, or results outside the approved scope of work.",
            ],
            True,
        ),
        (
            "Cancellation and Termination",
            [
                "Either party may cancel or terminate the services before completion by giving notice to the other party. If the Client cancels after work has begun, the Client is responsible for payment for all work performed, materials used, and costs incurred up to the time of cancellation. The Business may cancel, reschedule, or terminate the services if conditions are unsafe, access is not provided, required utilities are unavailable, payment is not made as agreed, or the Client prevents the Business from performing the work."
            ],
            True,
        ),
        (
            "Modifications",
            [
                "Any changes to the scope of work, price, schedule, or terms of this Contract must be agreed to by both the Client and the Business in writing. Written approval may include a signed change order, email, text message, Square message, or other written confirmation accepted by the Business."
            ],
            False,
        ),
        (
            "Licensing and Compliance",
            [
                "The Business will perform the services in a professional manner and in compliance with applicable laws and regulations. The Business is responsible for maintaining any licenses, permits, registrations, or insurance required for the services it provides, to the extent required by applicable law."
            ],
            False,
        ),
        (
            "Authority to Sign",
            [
                "Each party represents that it has the authority to enter into this Contract and to perform its obligations under this Contract. If the Client is signing on behalf of another person, business, property owner, or organization, the Client represents that they have authority to do so."
            ],
            False,
        ),
    ]

    for idx, (title_text, body, initials) in enumerate(clauses, 1):
        add_clause(doc, idx, title_text, body, initials)

    add_signature_block(doc)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Pressure Washing Service Agreement | Draft template")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(102, 112, 133)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
